# src/data_processing/resume_parser.py
"""Utilities for parsing candidate resumes (PDF or DOCX) into clean text."""
import io
import re
from typing import Any
import pdfplumber
import docx
from src.utils.logging import get_logger

logger = get_logger("data_processing.resume_parser")


def _extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF file given its raw bytes."""
    text_chunks = []
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page_idx, page in enumerate(pdf.pages, 1):
                try:
                    page_text = page.extract_text(layout=False)
                    if page_text and page_text.strip():
                        text_chunks.append(page_text.strip())
                except Exception as page_err:
                    logger.warning(f"Error extracting text from page {page_idx}: {page_err}")
    except Exception as e:
        logger.error(f"pdfplumber failed to open PDF: {e}")
        raise ValueError("Failed to parse PDF document. The file may be encrypted or corrupted.")

    return "\n\n".join(text_chunks)


def _extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX file given its raw bytes."""
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
        
        # Also extract table text from DOCX
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
                    
        return "\n".join(paragraphs)
    except Exception as e:
        logger.error(f"python-docx failed to open DOCX: {e}")
        raise ValueError("Failed to parse DOCX document. The file may be corrupted.")


def parse_resume(uploaded_file: Any) -> str:
    """Parse an uploaded resume file (PDF or DOCX) into clean plain text.

    Parameters
    ----------
    uploaded_file : UploadedFile
        The file object provided by Streamlit with ``name`` and ``read()``.

    Returns
    -------
    str
        Cleaned, normalized text representation of the resume.

    Raises
    ------
    ValueError
        If the file format is unsupported or text cannot be extracted.
    """
    if uploaded_file is None:
        raise ValueError("No file uploaded.")

    filename = getattr(uploaded_file, "name", "").lower()
    
    # Read bytes safely
    if hasattr(uploaded_file, "getvalue"):
        file_bytes = uploaded_file.getvalue()
    elif hasattr(uploaded_file, "read"):
        file_bytes = uploaded_file.read()
    else:
        file_bytes = bytes(uploaded_file)

    if not file_bytes:
        raise ValueError("Uploaded file is empty (0 bytes).")

    logger.info(f"Parsing uploaded resume: {filename} ({len(file_bytes)} bytes)")

    if filename.endswith(".pdf"):
        raw_text = _extract_text_from_pdf(file_bytes)
    elif filename.endswith(".docx") or filename.endswith(".doc"):
        raw_text = _extract_text_from_docx(file_bytes)
    else:
        raise ValueError("Unsupported file format. Please upload a PDF or DOCX file.")

    if not raw_text or len(raw_text.strip()) < 20:
        raise ValueError(
            "Unable to extract readable text from this file. "
            "Please ensure the document contains selectable text and is not an image-only scan."
        )

    # Normalize whitespace and redundant blank lines
    lines = [line.strip() for line in raw_text.splitlines() if line.strip()]
    cleaned = "\n".join(lines)

    logger.info(f"Successfully extracted {len(cleaned)} characters ({len(cleaned.split())} words) from resume.")
    return cleaned
